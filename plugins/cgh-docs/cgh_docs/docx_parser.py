# -#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#
# __creation__ = 2026-07-29
# __author__ = "jndjama (Joy Ndjama)"
# __copyright__ = "Copyright 2026 ALTIKVA."
# __licence__ = "MIT & CC BY-NC-SA (https://www.altikva.com/licenses/LICENSE-1.0)"
# -#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#
# Description: DOCX parser: Heading-styled paragraphs become the section
#              tree (mirroring the markdown parser), the body text between
#              headings becomes each section's searchable preview. "Lines"
#              are paragraph indexes, the closest stable anchor docx has.

from __future__ import annotations

import logging
import re
from pathlib import Path

from codegraph.plugin_api import BaseParser, FileIndex, SectionDef

_log = logging.getLogger(__name__)

_PREVIEW_CHARS = 400
_HEADING_RE = re.compile(r"heading\s*(\d)", re.IGNORECASE)


class DocxParser(BaseParser):
    """Best-effort docx to sections. A corrupt file yields an empty
    index, never an exception."""

    lang = "docx"
    extensions = [".docx"]
    extracts = ["sections"]
    description = "Word headings and body text as document sections"

    def parse(self, path: Path) -> FileIndex:
        idx = FileIndex(path=str(path), lang=self.lang)
        try:
            import docx

            document = docx.Document(str(path))
        except Exception:
            _log.warning("docx parse failed, indexed empty: %s", path)
            return idx

        current: SectionDef | None = None
        body: list[str] = []
        preamble: list[str] = []

        def _flush() -> None:
            if current is not None:
                current.body_preview = " ".join(body)[:_PREVIEW_CHARS]

        for i, para in enumerate(document.paragraphs, start=1):
            text = (para.text or "").strip()
            style = getattr(getattr(para, "style", None), "name", "") or ""
            m = _HEADING_RE.search(style)
            if m and text:
                _flush()
                if current is not None:
                    current.end_line = i - 1
                current = SectionDef(
                    id=f"{path}::para{i}",
                    title=text,
                    level=int(m.group(1)),
                    file_path=str(path),
                    start_line=i,
                    end_line=i,
                    anchor=f"para-{i}",
                )
                idx.sections.append(current)
                body = []
            elif text:
                (body if current is not None else preamble).append(text)

        _flush()
        if current is not None:
            current.end_line = len(document.paragraphs)

        # A document with no heading styles at all still gets one section
        # so its text is searchable.
        if not idx.sections and preamble:
            idx.sections.append(
                SectionDef(
                    id=f"{path}::body",
                    title=path.stem,
                    level=1,
                    file_path=str(path),
                    start_line=1,
                    end_line=len(document.paragraphs),
                    body_preview=" ".join(preamble)[:_PREVIEW_CHARS],
                    anchor="body",
                )
            )
        return idx
