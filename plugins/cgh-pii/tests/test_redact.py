# -#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#
# __creation__ = 2026-08-05
# __author__ = "jndjama (Joy Ndjama)"
# __copyright__ = "Copyright 2026 ALTIKVA."
# __licence__ = "MIT"
# -#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#
# Description: Redaction: regex categories (no NER needed), reading-order
#              numbering, stable tokens per value, keyed pseudonyms, the
#              category filter, and the clean error when a name is asked
#              for without the NER extra. The NER path itself is only
#              exercised when presidio is present.

from __future__ import annotations

import pytest

pytest.importorskip("cgh_pii")

from cgh_pii.redact import RedactError, redact


def test_only_filters_categories():
    t = "mail a@x.com, phone +33612345678"
    out, counts = redact(t, only=["email"])
    assert "[EMAIL_1]" in out
    assert "+33612345678" in out  # phone left alone
    assert counts == {"email": 1}


def test_placeholder_numbers_in_reading_order():
    out, _ = redact("first a@x.com then b@x.com", only=["email"])
    assert out == "first [EMAIL_1] then [EMAIL_2]"


def test_same_value_same_token():
    out, _ = redact("a@x.com and a@x.com again", only=["email"])
    assert out == "[EMAIL_1] and [EMAIL_1] again"


def test_pseudonym_is_keyed_and_stable():
    a, _ = redact("a@x.com", only=["email"], mode="pseudonym", secret=b"k" * 32)
    b, _ = redact("a@x.com", only=["email"], mode="pseudonym", secret=b"k" * 32)
    c, _ = redact("a@x.com", only=["email"], mode="pseudonym", secret=b"z" * 32)
    assert a == b and a != c
    assert a.startswith("<pii.email:") and "a@x.com" not in a


def test_iban_validated_not_any_digit_run():
    good = redact("IBAN FR7630006000011234567890189", only=["iban"])[0]
    assert "[IBAN_1]" in good
    bad = redact("ref FR0012345678901234567", only=["iban"])[0]
    assert "IBAN" not in bad  # fails mod-97, left as is


def test_unknown_category_raises():
    with pytest.raises(RedactError, match="unknown"):
        redact("x", only=["nope"])


def test_name_without_ner_raises_clearly(monkeypatch):
    import builtins

    real_import = builtins.__import__

    def no_presidio(name, *a, **k):
        if name.startswith("presidio"):
            raise ImportError("no presidio")
        return real_import(name, *a, **k)

    monkeypatch.setattr(builtins, "__import__", no_presidio)
    with pytest.raises(RedactError, match="NER"):
        redact("Jean Dupont", only=["person"])


def test_sdk_redact_text_delegates():
    from codegraph import sdk

    out = sdk.redact_text("mail a@x.com", only=["email"])
    assert out == "mail [EMAIL_1]"


@pytest.mark.network
class TestWithNer:
    """Only runs where presidio is installed."""

    def test_names_and_repeat_propagation(self):
        pytest.importorskip("presidio_analyzer")
        t = "Jean Dupont met Marie Curie. Jean Dupont left."
        out, counts = redact(t, only=["person"])
        assert out.count("[PERSON_1]") == 2  # both Jean Dupont
        assert "[PERSON_2]" in out  # Marie Curie
        assert counts["person"] == 3


class TestDocx:
    """python-docx is present in CI via cgh-docs; skip if absent."""

    def _doc(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document

        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Facture ")
        p.add_run("a@x.com").bold = True  # PII split across runs
        p.add_run(" et b@x.com.")
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = "Repondre a a@x.com"  # repeat, in a cell
        src = tmp_path / "in.docx"
        doc.save(str(src))
        return src

    def test_redacts_body_and_tables_with_shared_tokens(self, tmp_path):
        pytest.importorskip("docx")
        from cgh_pii.redact_docx import redact_docx_file
        from docx import Document

        src = self._doc(tmp_path)
        dst = tmp_path / "out.docx"
        counts = redact_docx_file(src, dst, only=["email"])
        out = Document(str(dst))
        body = out.paragraphs[0].text
        cell = out.tables[0].rows[0].cells[0].text
        assert "[EMAIL_1]" in body and "[EMAIL_2]" in body  # bold one caught
        assert "a@x.com" not in body and "a@x.com" not in cell
        # same value, same token across paragraph and table cell
        assert "[EMAIL_1]" in cell
        assert counts["email"] == 3

    def test_unchanged_paragraph_keeps_its_runs(self, tmp_path):
        pytest.importorskip("docx")
        from cgh_pii.redact_docx import redact_docx_file
        from docx import Document

        from_doc = Document()
        from_doc.add_paragraph("No PII here, just prose.")
        p = from_doc.add_paragraph()
        p.add_run("Bold").bold = True
        p.add_run(" and normal, no pii.")
        src = tmp_path / "clean.docx"
        from_doc.save(str(src))
        dst = tmp_path / "clean-out.docx"
        redact_docx_file(src, dst, only=["email"])
        out = Document(str(dst))
        # the mixed-formatting paragraph is untouched: still two runs
        assert len(out.paragraphs[1].runs) == 2
