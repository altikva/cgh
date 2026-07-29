# -#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#
# __creation__ = 2026-07-29
# __author__ = "jndjama (Joy Ndjama)"
# __copyright__ = "Copyright 2026 ALTIKVA."
# __licence__ = "MIT & CC BY-NC-SA (https://www.altikva.com/licenses/LICENSE-1.0)"
# -#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#
# Description: cgh-docs parser tests: pdf pages, docx heading tree, xlsx
#              sheets, corrupt-file resilience, and the end-to-end chain
#              (register through PluginAPI, index a repo, find document
#              sections through the FTS).

from __future__ import annotations

import subprocess

import pytest

pytest.importorskip("pypdf")
pytest.importorskip("docx")
pytest.importorskip("openpyxl")

import codegraph.plugins as plugins  # noqa: E402
from cgh_docs.docx_parser import DocxParser  # noqa: E402
from cgh_docs.pdf_parser import PdfParser  # noqa: E402
from cgh_docs.xlsx_parser import XlsxParser  # noqa: E402


def _mini_pdf() -> bytes:
    """A complete single-page PDF with one Helvetica text object, built
    with a correct xref table so pypdf accepts it. Handwritten because
    pypdf writes but does not typeset; extract_text yields
    "Hello cgh docs"."""
    stream = b"BT /F1 24 Tf 72 700 Td (Hello cgh docs) Tj ET"
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length %d >>\nstream\n%s\nendstream" % (len(stream), stream),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += b"%d 0 obj\n%s\nendobj\n" % (i, body)
    xref_at = len(out)
    out += b"xref\n0 %d\n" % (len(objects) + 1)
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += b"%010d 00000 n \n" % off
    out += b"trailer\n<< /Root 1 0 R /Size %d >>\nstartxref\n%d\n%%%%EOF\n" % (
        len(objects) + 1,
        xref_at,
    )
    return bytes(out)


def _make_docx(path):
    import docx

    d = docx.Document()
    d.add_heading("Contract terms", level=1)
    d.add_paragraph("The supplier delivers within thirty days.")
    d.add_heading("Payment", level=2)
    d.add_paragraph("Net sixty, in euros.")
    d.save(str(path))


def _make_xlsx(path):
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Donations"
    ws.append(["donor_email", "amount", "currency"])
    ws.append(["a@b.c", 100, "EUR"])
    wb.create_sheet("Refunds").append(["ref_id", "reason"])
    wb.save(str(path))


class TestPdfParser:
    def test_extracts_page_sections(self, tmp_path):
        f = tmp_path / "doc.pdf"
        f.write_bytes(_mini_pdf())
        idx = PdfParser().parse(f)
        assert idx.lang == "pdf"
        assert len(idx.sections) == 1
        assert idx.sections[0].title == "Page 1"
        assert "Hello cgh docs" in idx.sections[0].body_preview

    def test_corrupt_pdf_yields_empty_index(self, tmp_path):
        f = tmp_path / "bad.pdf"
        f.write_bytes(b"%PDF-not really")
        idx = PdfParser().parse(f)
        assert idx.sections == []


class TestDocxParser:
    def test_heading_tree_and_previews(self, tmp_path):
        f = tmp_path / "contract.docx"
        _make_docx(f)
        idx = DocxParser().parse(f)
        titles = [(s.title, s.level) for s in idx.sections]
        assert titles == [("Contract terms", 1), ("Payment", 2)]
        assert "thirty days" in idx.sections[0].body_preview
        assert "euros" in idx.sections[1].body_preview

    def test_corrupt_docx_yields_empty_index(self, tmp_path):
        f = tmp_path / "bad.docx"
        f.write_bytes(b"not a zip at all")
        assert DocxParser().parse(f).sections == []


class TestXlsxParser:
    def test_sheets_with_headers(self, tmp_path):
        f = tmp_path / "book.xlsx"
        _make_xlsx(f)
        idx = XlsxParser().parse(f)
        assert [s.title for s in idx.sections] == ["Donations", "Refunds"]
        assert "donor_email" in idx.sections[0].body_preview

    def test_corrupt_xlsx_yields_empty_index(self, tmp_path):
        f = tmp_path / "bad.xlsx"
        f.write_bytes(b"nope")
        assert XlsxParser().parse(f).sections == []


class TestEndToEnd:
    @pytest.fixture(autouse=True)
    def clean_registries(self):
        plugins._reset_for_tests()
        yield
        plugins._reset_for_tests()
        import codegraph.parsers as parsers

        for ext in (".pdf", ".docx", ".xlsx", ".xlsm"):
            parsers._REGISTRY.pop(ext, None)
            parsers._INSTANCES.pop(ext, None)

    def test_indexed_docx_is_searchable(self, tmp_path):
        import cgh_docs
        from codegraph.plugin_api import PluginAPI

        api = PluginAPI("docs", tmp_path, {}, plugins._registries)
        cgh_docs.register(api)

        subprocess.run(["git", "init", "-q"], cwd=tmp_path, check=True)
        _make_docx(tmp_path / "contract.docx")

        from codegraph.core.db import reset_connection
        from codegraph.indexer import index_repo

        reset_connection()
        try:
            stats = index_repo(str(tmp_path))
        finally:
            reset_connection()
        assert stats["errors"] == 0

        from codegraph.core.fts import fts_search, get_fts_conn

        hits = fts_search(get_fts_conn(tmp_path), "payment", limit=10)
        assert any(h.kind == "md_section" for h in hits)
