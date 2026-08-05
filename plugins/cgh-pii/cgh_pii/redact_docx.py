# -#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#
# __creation__ = 2026-08-05
# __author__ = "jndjama (Joy Ndjama)"
# __copyright__ = "Copyright 2026 ALTIKVA."
# __licence__ = "MIT"
# -#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#-#
# Description: docx redaction (cgh-pii[docx]). Redacts body paragraphs
#              and table cells with one shared token map, so the same
#              value maps to the same token across the whole document.
#              A changed paragraph is rewritten by collapsing its runs,
#              which flattens formatting inside that paragraph but is the
#              only way to redact PII that spans runs (a name split by a
#              bold surname); unchanged paragraphs keep their formatting.
#              Needs python-docx (the [docx] extra).

from __future__ import annotations

from pathlib import Path

from .redact import redact_chunks


def _paragraphs(doc):
    """Every editable paragraph: body, then table cells, in reading
    order. Headers and footers are out of scope for now."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _rewrite(paragraph, new_text: str) -> None:
    """Set the paragraph text, collapsing its runs. Keeps the first
    run (and its style) as the carrier so the paragraph does not lose
    its base formatting entirely."""
    runs = paragraph.runs
    if runs:
        runs[0].text = new_text
        for extra in runs[1:]:
            extra.text = ""
    else:
        paragraph.add_run(new_text)


def redact_docx_file(
    src: Path,
    dst: Path,
    only: list[str] | None = None,
    mode: str = "placeholder",
    secret: bytes | None = None,
    language: str = "en",
) -> dict[str, int]:
    """Redact ``src`` into ``dst``. Returns counts by category. Raises
    RedactError (via redact_chunks) if a requested NER category has no
    presidio. ImportError propagates if python-docx is not installed."""
    from docx import Document

    doc = Document(str(src))
    paras = list(_paragraphs(doc))
    texts = [p.text for p in paras]
    redacted, counts = redact_chunks(texts, only, mode, secret, language)
    for para, before, after in zip(paras, texts, redacted, strict=True):
        if after != before:
            _rewrite(para, after)
    doc.save(str(dst))
    return counts
